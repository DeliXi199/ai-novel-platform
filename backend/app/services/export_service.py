from __future__ import annotations

import re
from datetime import UTC, datetime
from io import BytesIO

from fastapi import HTTPException
from sqlalchemy.orm import Session

from app.models.chapter import Chapter
from app.models.novel import Novel


def _safe_filename(value: str) -> str:
    value = value.strip()
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value)
    value = re.sub(r"\s+", "_", value)
    return value[:120] or "novel"


def _build_metadata_lines(novel: Novel) -> list[str]:
    lines = [
        f"书名：{novel.title}",
        f"题材：{novel.genre}",
        f"主角：{novel.protagonist_name}",
        f"简介：{novel.premise}",
    ]

    tone = None
    forbidden = None
    if isinstance(novel.style_preferences, dict):
        tone = novel.style_preferences.get("tone")
        forbidden = novel.style_preferences.get("forbidden")

    if tone:
        lines.append(f"风格：{tone}")
    if forbidden:
        if isinstance(forbidden, list):
            forbidden_text = "、".join(str(item) for item in forbidden)
        else:
            forbidden_text = str(forbidden)
        lines.append(f"限制：{forbidden_text}")

    lines.append(f"当前章节数：{novel.current_chapter_no}")
    lines.append(f"导出时间：{datetime.now(UTC).strftime('%Y-%m-%d %H:%M:%S')} UTC")
    return lines


def render_novel_as_txt(novel: Novel, chapters: list[Chapter]) -> str:
    metadata = _build_metadata_lines(novel)
    parts: list[str] = ["\n".join(metadata), "", "=" * 48, ""]

    for chapter in chapters:
        parts.append(f"第{chapter.chapter_no}章 {chapter.title}")
        parts.append("")
        parts.append(chapter.content.strip())
        parts.append("")
        parts.append("-" * 48)
        parts.append("")

    return "\n".join(parts).strip() + "\n"


def render_novel_as_markdown(novel: Novel, chapters: list[Chapter]) -> str:
    metadata = _build_metadata_lines(novel)
    parts: list[str] = [f"# {novel.title}", ""]
    parts.extend([f"- {line}" for line in metadata[1:]])
    parts.append("")
    parts.append("---")
    parts.append("")

    for chapter in chapters:
        parts.append(f"## 第{chapter.chapter_no}章 {chapter.title}")
        parts.append("")
        content = chapter.content.strip().replace("\r\n", "\n")
        parts.append(content)
        parts.append("")

    return "\n".join(parts).strip() + "\n"


def render_novel_as_docx(novel: Novel, chapters: list[Chapter]) -> BytesIO:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(novel.title)
    run.bold = True
    run.font.size = Pt(18)

    for line in _build_metadata_lines(novel)[1:]:
        p = doc.add_paragraph()
        p.add_run(line)

    doc.add_page_break()

    for idx, chapter in enumerate(chapters, start=1):
        if idx > 1:
            doc.add_page_break()
        h = doc.add_paragraph()
        h_run = h.add_run(f"第{chapter.chapter_no}章 {chapter.title}")
        h_run.bold = True
        h_run.font.size = Pt(15)

        for para in chapter.content.strip().replace("\r\n", "\n").split("\n"):
            if not para.strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.line_spacing = 1.5
            p.add_run(para.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def render_novel_as_pdf(novel: Novel, chapters: list[Chapter]) -> BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "NovelTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=20,
        leading=24,
        alignment=1,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "NovelMeta",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=16,
        spaceAfter=4,
    )
    chapter_title_style = ParagraphStyle(
        "ChapterTitle",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=16,
        leading=20,
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=11.5,
        leading=19,
        firstLineIndent=22,
        spaceAfter=6,
    )

    story: list = []
    story.append(Paragraph(novel.title, title_style))
    for line in _build_metadata_lines(novel)[1:]:
        story.append(Paragraph(line, meta_style))
    story.append(PageBreak())

    for idx, chapter in enumerate(chapters, start=1):
        if idx > 1:
            story.append(PageBreak())
        story.append(Paragraph(f"第{chapter.chapter_no}章 {chapter.title}", chapter_title_style))
        for para in chapter.content.strip().replace("\r\n", "\n").split("\n"):
            para = para.strip()
            if not para:
                continue
            para = (
                para.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(para, body_style))
            story.append(Spacer(1, 2))

    def add_page_number(canvas, _doc):
        canvas.setFont("STSong-Light", 9)
        canvas.drawCentredString(A4[0] / 2, 10 * mm, str(canvas.getPageNumber()))

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    buffer.seek(0)
    return buffer


def export_novel_bytes(db: Session, novel_id: int, export_format: str) -> tuple[BytesIO, str, str]:
    novel = db.query(Novel).filter(Novel.id == novel_id).first()
    if not novel:
        raise HTTPException(status_code=404, detail="Novel not found")

    chapters = (
        db.query(Chapter)
        .filter(Chapter.novel_id == novel_id)
        .order_by(Chapter.chapter_no.asc())
        .all()
    )
    if not chapters:
        raise HTTPException(status_code=404, detail="No chapters found for export")

    export_format = export_format.lower().strip()
    if export_format not in {"txt", "md", "docx", "pdf"}:
        raise HTTPException(status_code=400, detail="Unsupported export format. Use 'txt', 'md', 'docx' or 'pdf'.")

    if export_format == "txt":
        content = render_novel_as_txt(novel, chapters)
        media_type = "text/plain; charset=utf-8"
        extension = "txt"
        buffer = BytesIO(content.encode("utf-8"))
    elif export_format == "md":
        content = render_novel_as_markdown(novel, chapters)
        media_type = "text/markdown; charset=utf-8"
        extension = "md"
        buffer = BytesIO(content.encode("utf-8"))
    elif export_format == "docx":
        buffer = render_novel_as_docx(novel, chapters)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        extension = "docx"
    else:
        buffer = render_novel_as_pdf(novel, chapters)
        media_type = "application/pdf"
        extension = "pdf"

    filename = f"{_safe_filename(novel.title)}.{extension}"
    return buffer, filename, media_type
